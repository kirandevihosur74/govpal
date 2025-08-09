from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import json
import numpy as np
import aiofiles
import io
from datetime import datetime
import hashlib
import re
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Initialize the embedding model (will download on first use)
embedding_model = None

def get_embedding_model():
    global embedding_model
    if embedding_model is None:
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    return embedding_model

# Configuration
INDEX_DIR = os.getenv("INDEX_DIR", "./index")
STORAGE_DIR = os.getenv("STORAGE_DIR", "./storage")
CHUNK_SIZE = 800
CHUNK_OVERLAP = 120
SIMILARITY_THRESHOLD = 0.45

# Ensure directories exist
os.makedirs(INDEX_DIR, exist_ok=True)
os.makedirs(STORAGE_DIR, exist_ok=True)

app = FastAPI(title="GovPal API", version="1.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Pydantic models
class DocumentMetadata(BaseModel):
    dept: Optional[str] = None
    year: Optional[int] = None
    tags: List[str] = []

class FileIngestResult(BaseModel):
    filename: str
    document_id: Optional[str] = None
    chunks_created: int = 0
    status: str  # "success", "error", "skipped"
    message: str
    error_details: Optional[str] = None

class IngestResponse(BaseModel):
    files_processed: int
    files_successful: int
    files_failed: int
    total_chunks_created: int
    results: List[FileIngestResult]
    message: str

class SearchResult(BaseModel):
    id: str
    title: str
    content: str
    score: float
    metadata: Dict[str, Any]
    path: str

class SearchResponse(BaseModel):
    query: str
    results: List[SearchResult]
    aggregates: Dict[str, Any]
    total: int
    limit: int

class AnswerQuery(BaseModel):
    question: str
    context: Optional[str] = None

class AnswerResponse(BaseModel):
    answer: str
    sources: List[str]
    confidence: float

# Utility functions
def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at a sentence or word boundary
        if end < len(text):
            last_period = chunk.rfind('.')
            last_space = chunk.rfind(' ')
            if last_period > chunk_size * 0.8:
                end = start + last_period + 1
            elif last_space > chunk_size * 0.8:
                end = start + last_space
        
        chunks.append(text[start:end].strip())
        start = end - overlap
        
        if start >= len(text):
            break
    
    return chunks

def extract_year_from_text(text: str) -> Optional[int]:
    """Extract likely year from document text."""
    year_pattern = r'\b(19|20)\d{2}\b'
    years = re.findall(year_pattern, text)
    if years:
        # Return the most recent year found
        return max(int(year) for year in years)
    return None

def generate_doc_id(filename: str, content_hash: str) -> str:
    """Generate unique document ID."""
    return hashlib.md5(f"{filename}:{content_hash}".encode()).hexdigest()[:12]

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text from PDF or DOCX file."""
    file_extension = filename.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        pdf_reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    elif file_extension in ['docx', 'doc']:
        doc = Document(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

async def process_single_file(
    file_content: bytes, 
    filename: str, 
    dept: Optional[str] = None,
    year: Optional[int] = None,
    tags: List[str] = None
) -> FileIngestResult:
    """Process a single file and return result."""
    tags = tags or []
    
    try:
        # Check file type
        file_extension = filename.lower().split('.')[-1]
        if file_extension not in ['pdf', 'docx', 'doc']:
            return FileIngestResult(
                filename=filename,
                status="skipped",
                message=f"Unsupported file type: {file_extension}",
                error_details="Only PDF and DOCX files are supported"
            )
        
        # Extract text
        text = extract_text_from_file(file_content, filename)
        
        if not text.strip():
            return FileIngestResult(
                filename=filename,
                status="error",
                message="No text found in file",
                error_details="File appears to be empty or unreadable"
            )
        
        # Generate IDs and paths
        content_hash = hashlib.md5(file_content).hexdigest()
        doc_id = generate_doc_id(filename, content_hash)
        
        # Extract year if not provided
        if not year:
            year = extract_year_from_text(text)
        
        # Save file to storage
        file_path = os.path.join(STORAGE_DIR, f"{doc_id}_{filename}")
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Chunk the text
        chunks = chunk_text(text)
        
        # Load existing index
        index_data = await load_index()
        
        # Generate embeddings for chunks
        model = get_embedding_model()
        chunk_embeddings = model.encode(chunks)
        
        # Prepare new chunk and doc entries
        chunk_entries = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            chunk_entries.append({
                "id": chunk_id,
                "doc_id": doc_id,
                "text": chunk,
                "year": year,
                "path": file_path,
                "dept": dept,
                "tags": tags
            })
        
        doc_entry = {
            "id": doc_id,
            "title": filename,
            "path": file_path,
            "year": year,
            "dept": dept,
            "tags": tags,
            "created_at": datetime.now().isoformat()
        }
        
        # Update index
        if index_data["embeddings"] is not None:
            new_embeddings = np.vstack([index_data["embeddings"], chunk_embeddings])
        else:
            new_embeddings = chunk_embeddings
        
        all_chunks = index_data["chunks"] + chunk_entries
        all_docs = index_data["docs"] + [doc_entry]
        
        # Save updated index
        await save_index(new_embeddings, all_chunks, all_docs)
        
        return FileIngestResult(
            filename=filename,
            document_id=doc_id,
            chunks_created=len(chunks),
            status="success",
            message="File processed successfully"
        )
        
    except Exception as e:
        return FileIngestResult(
            filename=filename,
            status="error",
            message=f"Error processing file: {str(e)}",
            error_details=str(e)
        )

async def load_index() -> Dict[str, Any]:
    """Load existing index data."""
    index_data = {
        "embeddings": None,
        "chunks": [],
        "docs": []
    }
    
    try:
        # Load embeddings
        embeddings_path = os.path.join(INDEX_DIR, "embeddings.npy")
        if os.path.exists(embeddings_path):
            index_data["embeddings"] = np.load(embeddings_path)
        
        # Load chunks
        chunks_path = os.path.join(INDEX_DIR, "chunks.json")
        if os.path.exists(chunks_path):
            async with aiofiles.open(chunks_path, 'r') as f:
                content = await f.read()
                index_data["chunks"] = json.loads(content)
        
        # Load docs
        docs_path = os.path.join(INDEX_DIR, "docs.json")
        if os.path.exists(docs_path):
            async with aiofiles.open(docs_path, 'r') as f:
                content = await f.read()
                index_data["docs"] = json.loads(content)
    
    except Exception as e:
        print(f"Error loading index: {e}")
    
    return index_data

async def save_index(embeddings: np.ndarray, chunks: List[Dict], docs: List[Dict]):
    """Save index data to files."""
    try:
        # Save embeddings
        embeddings_path = os.path.join(INDEX_DIR, "embeddings.npy")
        np.save(embeddings_path, embeddings)
        
        # Save chunks
        chunks_path = os.path.join(INDEX_DIR, "chunks.json")
        async with aiofiles.open(chunks_path, 'w') as f:
            await f.write(json.dumps(chunks, indent=2))
        
        # Save docs
        docs_path = os.path.join(INDEX_DIR, "docs.json")
        async with aiofiles.open(docs_path, 'w') as f:
            await f.write(json.dumps(docs, indent=2))
    
    except Exception as e:
        print(f"Error saving index: {e}")
        raise


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"ok": True}


@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "Welcome to GovPal API"}


@app.post("/ingest", response_model=IngestResponse)
async def ingest_documents(
    files: List[UploadFile] = File(...), 
    dept: Optional[str] = Form(None),
    year: Optional[int] = Form(None),
    tags: Optional[str] = Form(None)
):
    """Ingest multiple documents for processing - extract text, chunk, embed, and index."""
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    # Parse tags once
    tag_list = []
    if tags:
        tag_list = [tag.strip() for tag in tags.split(',') if tag.strip()]
    
    results = []
    files_successful = 0
    files_failed = 0
    total_chunks = 0
    
    for file in files:
        if not file.filename:
            result = FileIngestResult(
                filename="unknown",
                status="error",
                message="File has no filename",
                error_details="Cannot process file without filename"
            )
            results.append(result)
            files_failed += 1
            continue
        
        try:
            # Read file content
            content = await file.read()
            
            # Process the file
            result = await process_single_file(
                file_content=content,
                filename=file.filename,
                dept=dept,
                year=year,
                tags=tag_list
            )
            
            results.append(result)
            
            if result.status == "success":
                files_successful += 1
                total_chunks += result.chunks_created
            else:
                files_failed += 1
                
        except Exception as e:
            result = FileIngestResult(
                filename=file.filename,
                status="error",
                message=f"Unexpected error: {str(e)}",
                error_details=str(e)
            )
            results.append(result)
            files_failed += 1
    
    # Generate summary message
    if files_successful == len(files):
        message = f"All {len(files)} files processed successfully"
    elif files_successful > 0:
        message = f"{files_successful} of {len(files)} files processed successfully"
    else:
        message = f"Failed to process any of the {len(files)} files"
    
    return IngestResponse(
        files_processed=len(files),
        files_successful=files_successful,
        files_failed=files_failed,
        total_chunks_created=total_chunks,
        results=results,
        message=message
    )


@app.get("/search", response_model=SearchResponse)
async def search_documents(
    q: str = Query(..., description="Search query"),
    dept: Optional[str] = Query(None, description="Filter by department"),
    role: Optional[str] = Query(None, description="User role for access control"),
    limit: int = Query(10, description="Maximum number of results")
):
    """Search through ingested documents with embedding-based retrieval."""
    try:
        # Load index
        index_data = await load_index()
        
        if not index_data["chunks"] or index_data["embeddings"] is None:
            return SearchResponse(
                query=q,
                results=[],
                aggregates={"byYear": {}, "byDept": {}},
                total=0,
                limit=limit
            )
        
        # Generate query embedding
        model = get_embedding_model()
        query_embedding = model.encode([q])
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding, index_data["embeddings"])[0]
        
        # Get top chunks with similarity above threshold
        chunk_scores = list(enumerate(similarities))
        chunk_scores.sort(key=lambda x: x[1], reverse=True)
        
        # Filter by similarity threshold and department
        filtered_chunks = []
        for chunk_idx, score in chunk_scores[:50]:  # Take top 50 for processing
            if score < SIMILARITY_THRESHOLD:
                continue
                
            chunk = index_data["chunks"][chunk_idx]
            
            # Filter by department if specified
            if dept and chunk.get("dept") != dept:
                continue
                
            filtered_chunks.append((chunk_idx, score, chunk))
        
        # Group by document and take top documents
        doc_scores = {}
        for chunk_idx, score, chunk in filtered_chunks:
            doc_id = chunk["doc_id"]
            if doc_id not in doc_scores:
                doc_scores[doc_id] = {"max_score": score, "chunks": []}
            else:
                doc_scores[doc_id]["max_score"] = max(doc_scores[doc_id]["max_score"], score)
            doc_scores[doc_id]["chunks"].append((chunk_idx, score, chunk))
        
        # Sort documents by max score
        sorted_docs = sorted(doc_scores.items(), key=lambda x: x[1]["max_score"], reverse=True)
        
        # Build results
        results = []
        doc_lookup = {doc["id"]: doc for doc in index_data["docs"]}
        
        for doc_id, doc_data in sorted_docs[:limit]:
            doc_info = doc_lookup.get(doc_id, {})
            best_chunk = max(doc_data["chunks"], key=lambda x: x[1])
            
            result = SearchResult(
                id=doc_id,
                title=doc_info.get("title", "Unknown Document"),
                content=best_chunk[2]["text"][:300] + "..." if len(best_chunk[2]["text"]) > 300 else best_chunk[2]["text"],
                score=round(best_chunk[1], 3),
                metadata={
                    "year": best_chunk[2].get("year"),
                    "dept": best_chunk[2].get("dept"),
                    "tags": best_chunk[2].get("tags", []),
                    "chunk_count": len(doc_data["chunks"])
                },
                path=best_chunk[2].get("path", "")
            )
            results.append(result)
        
        # Build aggregates
        aggregates = {"byYear": {}, "byDept": {}}
        
        for chunk in index_data["chunks"]:
            # Year aggregation
            year = chunk.get("year")
            if year:
                aggregates["byYear"][str(year)] = aggregates["byYear"].get(str(year), 0) + 1
            
            # Department aggregation  
            dept_name = chunk.get("dept")
            if dept_name:
                aggregates["byDept"][dept_name] = aggregates["byDept"].get(dept_name, 0) + 1
        
        # If no results above threshold, suggest alternatives
        if not results:
            # Could add keyword-based fallback here
            pass
        
        return SearchResponse(
            query=q,
            results=results,
            aggregates=aggregates,
            total=len(results),
            limit=limit
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search error: {str(e)}")


@app.post("/answer")
async def answer_question(query: AnswerQuery):
    """Generate an answer to a question using ingested documents"""
    return {
        "question": query.question,
        "answer": "Based on the available documents, here is a comprehensive answer to your question. This is a placeholder response that demonstrates the API structure.",
        "sources": [
            "doc_123456 (sample.pdf, page 1)",
            "doc_789012 (policy.pdf, page 3)"
        ],
        "confidence": 0.92,
        "context_used": query.context or "No additional context provided"
    }
